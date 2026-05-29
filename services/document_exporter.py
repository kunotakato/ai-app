from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics


def export_to_docx(content: str) -> bytes:
    """
    生成結果をWordファイルに変換する関数です。
    戻り値はStreamlitのdownload_buttonで使えるbytesです。
    """

    document = Document()

    document.add_heading("職務経歴書 下書き", level=1)

    for line in content.split("\n"):
        line = line.strip()

        if not line:
            document.add_paragraph("")
        elif line.startswith("# "):
            document.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            document.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("### "):
            document.add_heading(line.replace("### ", ""), level=3)
        elif line.startswith("- ") or line.startswith("・"):
            document.add_paragraph(line, style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def export_to_pdf(content: str) -> bytes:
    """
    生成結果をPDFファイルに変換する関数です。
    日本語文字化け対策として HeiseiKakuGo-W5 を使用します。
    """

    buffer = BytesIO()

    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))

    c = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4
    margin_x = 40
    margin_y = 50
    line_height = 16

    x = margin_x
    y = height - margin_y

    c.setFont("HeiseiKakuGo-W5", 11)

    for raw_line in content.split("\n"):
        line = raw_line.strip()

        if not line:
            y -= line_height
            continue

        if line.startswith("# "):
            c.setFont("HeiseiKakuGo-W5", 16)
            line = line.replace("# ", "")
        elif line.startswith("## "):
            c.setFont("HeiseiKakuGo-W5", 13)
            line = line.replace("## ", "")
        elif line.startswith("### "):
            c.setFont("HeiseiKakuGo-W5", 12)
            line = line.replace("### ", "")
        else:
            c.setFont("HeiseiKakuGo-W5", 10)

        wrapped_lines = _wrap_text(line, max_chars=48)

        for wrapped_line in wrapped_lines:
            if y < margin_y:
                c.showPage()
                c.setFont("HeiseiKakuGo-W5", 10)
                y = height - margin_y

            c.drawString(x, y, wrapped_line)
            y -= line_height

    c.save()
    buffer.seek(0)

    return buffer.getvalue()


def _wrap_text(text: str, max_chars: int = 48) -> list[str]:
    """
    PDFで長い文章が右にはみ出さないように、
    指定文字数ごとに折り返します。
    """

    return [text[i:i + max_chars] for i in range(0, len(text), max_chars)]